"""Shared fixture builders: synthetic emails, MSG/OLE files, DOCX, PDFs.

Everything is generated in pytest temporary directories; no client files,
filenames, or excerpts are used.
"""

import datetime
import io
import struct
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

_FILETIME_EPOCH = datetime.datetime(1601, 1, 1, tzinfo=datetime.timezone.utc)


def to_filetime(dt: datetime.datetime) -> int:
    return int((dt - _FILETIME_EPOCH).total_seconds() * 10_000_000)


# ---------------------------------------------------------------------------
# .eml / reference renderer (copy of the pre-migration implementation)
# ---------------------------------------------------------------------------


def reference_eml_markdown(source: Path) -> str:
    """The exact pre-migration _convert_one_email rendering, as a reference."""
    import email as _email
    from email import policy as _policy

    with open(source, "rb") as f:
        msg = _email.message_from_binary_file(f, policy=_policy.default)

    lines = [
        f"# {msg['subject'] or '(no subject)'}",
        "",
        f"**From:** {msg['from']}",
        f"**To:** {msg['to']}",
    ]
    if msg["cc"]:
        lines.append(f"**CC:** {msg['cc']}")
    lines += [f"**Date:** {msg['date']}", "", "---", ""]

    body = None
    for part in msg.walk():
        if part.get_content_type() == "text/plain":
            try:
                body = part.get_content()
            except Exception:
                body = part.get_payload(decode=True).decode("utf-8", errors="replace")
            break
    lines.append(body if body else "(no text content)")
    return "\n".join(lines)


SIMPLE_EML = (
    b"From: Alice Author <alice@example.com>\r\n"
    b"To: Bob Reader <bob@example.com>\r\n"
    b"Subject: Simple message\r\n"
    b"Date: Tue, 04 Aug 2026 10:00:00 -0500\r\n"
    b"\r\n"
    b"First line.\r\nSecond line.\r\n"
)

MULTIPART_EML = (
    b"From: alice@example.com\r\n"
    b"To: bob@example.com\r\n"
    b"CC: carol@example.com\r\n"
    b"Subject: =?utf-8?B?UsOpc3Vtw6kgcmV2aWV3?=\r\n"
    b"Date: Tue, 04 Aug 2026 11:00:00 -0500\r\n"
    b"MIME-Version: 1.0\r\n"
    b"Content-Type: multipart/alternative; boundary=BOUND\r\n"
    b"\r\n"
    b"--BOUND\r\n"
    b"Content-Type: text/plain; charset=utf-8\r\n"
    b"Content-Transfer-Encoding: base64\r\n"
    b"\r\n"
    b"Q2Fmw6kgbcOpbnUgYXR0YWNoZWQu\r\n"
    b"--BOUND\r\n"
    b"Content-Type: text/html; charset=utf-8\r\n"
    b"\r\n"
    b"<b>Caf\xc3\xa9 m\xc3\xa9nu attached.</b>\r\n"
    b"--BOUND--\r\n"
)

HTML_ONLY_EML = (
    b"From: alice@example.com\r\n"
    b"To: bob@example.com\r\n"
    b"Subject: html only\r\n"
    b"Date: Tue, 04 Aug 2026 12:00:00 -0500\r\n"
    b"Content-Type: text/html\r\n"
    b"\r\n"
    b"<p>No plain text part here.</p>\r\n"
)

BAD_CHARSET_EML = (
    b"From: alice@example.com\r\n"
    b"To: bob@example.com\r\n"
    b"Subject: bad charset\r\n"
    b"Date: Tue, 04 Aug 2026 13:00:00 -0500\r\n"
    b"Content-Type: text/plain; charset=not-a-charset\r\n"
    b"\r\n"
    b"body with raw bytes \xff\xfe end\r\n"
)


def make_mht(html_part: bytes | None, text_part: bytes | None) -> bytes:
    parts = []
    if text_part is not None:
        parts.append(
            b"--MHTBOUND\r\nContent-Type: text/plain; charset=utf-8\r\n\r\n"
            + text_part
            + b"\r\n"
        )
    if html_part is not None:
        parts.append(
            b"--MHTBOUND\r\nContent-Type: text/html; charset=utf-8\r\n\r\n"
            + html_part
            + b"\r\n"
        )
    return (
        b"From: <Saved by Test>\r\n"
        b"Subject: Saved Page\r\n"
        b"Date: Tue, 04 Aug 2026 14:00:00 -0500\r\n"
        b"MIME-Version: 1.0\r\n"
        b"Content-Type: multipart/related; boundary=MHTBOUND\r\n"
        b"\r\n" + b"".join(parts) + b"--MHTBOUND--\r\n"
    )


def make_mbox(messages: list[bytes]) -> bytes:
    out = []
    for raw in messages:
        out.append(b"From alice@example.com Tue Aug  4 10:00:00 2026\n")
        out.append(raw.replace(b"\r\n", b"\n"))
        if not raw.endswith(b"\n"):
            out.append(b"\n")
        out.append(b"\n")
    return b"".join(out)


# ---------------------------------------------------------------------------
# Synthetic Outlook .msg / .oft (OLE) builder
# ---------------------------------------------------------------------------


def make_msg(
    path: Path,
    *,
    subject: str | None = None,
    sender: str | None = None,
    recipients: list[tuple[str, str, int]] = (),  # (name, email, 1=to 2=cc)
    body: str | None = None,
    html_body: bytes | None = None,
    compressed_rtf: bytes | None = None,
    ansi: bool = False,
    date: datetime.datetime | None = None,
    attachment_names: list[str] = (),
) -> Path:
    """Write a minimal but valid MSG/OFT OLE file for extract-msg."""
    from extract_msg.ole_writer import OleWriter

    writer = OleWriter()
    string_type = "001E" if ansi else "001F"

    def encode(value: str) -> bytes:
        return value.encode("cp1252") if ansi else value.encode("utf-16-le")

    def add_string(pid: str, value: str | None, prefix: str = "") -> None:
        if value is not None:
            writer.addEntry(f"{prefix}__substg1.0_{pid}{string_type}", encode(value))

    def add_binary(pid: str, value: bytes | None, prefix: str = "") -> None:
        if value is not None:
            writer.addEntry(f"{prefix}__substg1.0_{pid}0102", value)

    # Named-properties storage (required by extract-msg when attachments
    # or named props are accessed); empty streams satisfy the standard.
    writer.addEntry("__nameid_version1.0/__substg1.0_00020102", b"")
    writer.addEntry("__nameid_version1.0/__substg1.0_00030102", b"")
    writer.addEntry("__nameid_version1.0/__substg1.0_00040102", b"")

    add_string("001A", "IPM.Note")
    add_string("0037", subject)
    add_string("0C1A", sender)
    add_string("1000", body)
    add_binary("1013", html_body)
    add_binary("1009", compressed_rtf)

    for i, (name, email_addr, rtype) in enumerate(recipients):
        prefix = f"__recip_version1.0_#{i:08X}/"
        add_string("3001", name, prefix)
        add_string("39FE", email_addr, prefix)
        add_string("3003", email_addr, prefix)
        props = struct.pack("<IIIi", 0x0C150003, 0x6, rtype, 0)
        props += struct.pack("<IIIi", 0x30000003, 0x6, i, 0)
        writer.addEntry(prefix + "__properties_version1.0", b"\x00" * 8 + props)

    for i, name in enumerate(attachment_names):
        prefix = f"__attach_version1.0_#{i:08X}/"
        add_string("3707", name, prefix)
        add_string("3704", name, prefix)
        add_binary("3701", b"attachment-bytes", prefix)
        props = struct.pack("<IIIi", 0x37050003, 0x6, 1, 0)  # ATTACH_BY_VALUE
        writer.addEntry(prefix + "__properties_version1.0", b"\x00" * 8 + props)

    top_entries = b""
    if date is not None:
        top_entries += struct.pack("<IIQ", 0x00390040, 0x6, to_filetime(date))
    header = struct.pack(
        "<8sIIII8s",
        b"\x00" * 8,
        len(recipients),
        len(attachment_names),
        len(recipients),
        len(attachment_names),
        b"\x00" * 8,
    )
    writer.addEntry("__properties_version1.0", header + top_entries)
    writer.write(str(path))
    return path


# ---------------------------------------------------------------------------
# DOCX / PDF builders
# ---------------------------------------------------------------------------


def tiny_png_bytes() -> bytes:
    import fitz

    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 8, 8), False)
    pix.clear_with(128)
    return pix.tobytes("png")


def make_docx(
    path: Path,
    *,
    heading: str = "Agreement Title",
    paragraphs: list[str] = ("Introductory paragraph one.",),
    nested_list: bool = True,
    table: list[list[str]] | None = None,
    image: bool = False,
) -> Path:
    import docx

    d = docx.Document()
    d.add_heading(heading, level=1)
    for text in paragraphs:
        d.add_paragraph(text)
    if nested_list:
        d.add_paragraph("First numbered item", style="List Number")
        d.add_paragraph("Nested numbered item", style="List Number 2")
    if table is not None:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    if image:
        d.add_picture(io.BytesIO(tiny_png_bytes()))
    d.save(str(path))
    return path


def make_digital_pdf(path: Path, texts: list[str] = ("Digital page text.",)) -> Path:
    import fitz

    doc = fitz.open()
    for text in texts:
        page = doc.new_page()
        y = 100
        for line in text.split("\n"):
            page.insert_text((72, y), line, fontsize=11)
            y += 16
    doc.save(str(path))
    doc.close()
    return path


def make_scanned_pdf(path: Path, pages: int = 2) -> Path:
    import fitz

    doc = fitz.open()
    png = tiny_png_bytes()
    for _ in range(pages):
        page = doc.new_page()
        page.insert_image(page.rect, stream=png)
    doc.save(str(path))
    doc.close()
    return path


def make_mixed_pdf(path: Path) -> Path:
    """One digital page + two blank pages → majority blank → mixed/other."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Some digital text on the first page only.", fontsize=11)
    doc.new_page()
    doc.new_page()
    doc.save(str(path))
    doc.close()
    return path


def make_encrypted_pdf(path: Path) -> Path:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Secret text.", fontsize=11)
    doc.save(
        str(path),
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner-pw",
        user_pw="user-pw",
    )
    doc.close()
    return path


def make_malformed_pdf(path: Path) -> Path:
    path.write_bytes(b"%PDF-1.7\n" + bytes(range(256)) * 16 + b"\n%%EOF")
    return path


@pytest.fixture
def repo_tmp(tmp_path, monkeypatch):
    """A temporary working folder that startup.main() treats as the repo root.

    Hermetic: the repo-settings seam points at a tmp file so tests never read
    (or depend on) the real repo-root settings.json, and the sidecar-style
    global is restored after each test.
    """
    import repo_settings
    import startup

    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", ["startup.py"])
    monkeypatch.setattr(repo_settings, "SETTINGS_PATH", tmp_path / "settings.json")
    monkeypatch.setattr(startup, "SIDECAR_DOTFILES", False)
    monkeypatch.setattr(startup, "OCR_INT8", True)
    return tmp_path
